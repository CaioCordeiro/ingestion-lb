import io
import os
import sys
import uuid
from datetime import datetime

from fastapi import FastAPI, File, Header, HTTPException, UploadFile
from fastapi.responses import JSONResponse

# Add parent directory to path to import shared modules
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from shared.logging_config import configure_logger
from shared.rabbitmq import RabbitMQProducer
from shared.schemas import RawTextMessage

# Configure logger
logger = configure_logger("ingestion_file")

# Create FastAPI app
app = FastAPI(
    title="File Ingestion Service",
    description="Service for ingesting PDF and DOCX files",
    version="1.0.0",
)

# RabbitMQ configuration
RABBITMQ_HOST = os.getenv("RABBITMQ_HOST", "rabbitmq")
RABBITMQ_PORT = int(os.getenv("RABBITMQ_PORT", "5672"))
RABBITMQ_USER = os.getenv("RABBITMQ_USER", "guest")
RABBITMQ_PASS = os.getenv("RABBITMQ_PASS", "guest")
RABBITMQ_QUEUE = os.getenv("RABBITMQ_QUEUE", "raw_text_queue")

# Create RabbitMQ producer
rabbitmq_producer = RabbitMQProducer(
    host=RABBITMQ_HOST,
    port=RABBITMQ_PORT,
    username=RABBITMQ_USER,
    password=RABBITMQ_PASS,
    queue_name=RABBITMQ_QUEUE,
    logger=logger,
)


# Connect to RabbitMQ on startup
@app.on_event("startup")
async def startup_event():
    """Connect to RabbitMQ on startup."""
    rabbitmq_producer.connect()
    logger.info("File Ingestion Service started")


# Shutdown event
@app.on_event("shutdown")
async def shutdown_event():
    """Close connections on shutdown."""
    rabbitmq_producer.close()
    logger.info("File Ingestion Service shutting down")


async def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from a PDF file.

    Args:
        file_content: PDF file content as bytes

    Returns:
        Extracted text as string
    """
    try:
        from pypdf import PdfReader

        # Create a PDF reader object
        pdf_reader = PdfReader(io.BytesIO(file_content))

        # Extract text from each page
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"

        return text

    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise Exception(f"Error extracting text from PDF: {str(e)}")


async def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from a DOCX file.

    Args:
        file_content: DOCX file content as bytes

    Returns:
        Extracted text as string
    """
    try:
        from docx import Document

        # Create a Document object
        document = Document(io.BytesIO(file_content))

        # Extract text from paragraphs
        text = ""
        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"

        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"

        return text

    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise Exception(f"Error extracting text from DOCX: {str(e)}")


# Routes
@app.post("/ingest")
async def ingest_file(file: UploadFile = File(...), x_correlation_id: str = Header(None)):
    """
    Ingest PDF or DOCX file, extract text, and publish to RabbitMQ.

    Args:
        file: PDF or DOCX file to ingest
        x_correlation_id: Correlation ID from request header

    Returns:
        JSON response with status and correlation ID
    """
    # Use provided correlation ID or generate a new one
    correlation_id = x_correlation_id or str(uuid.uuid4())

    logger.info(
        f"Processing file ingestion request: {file.filename}",
        extra={"correlation_id": correlation_id},
    )

    # Check file extension
    file_ext = os.path.splitext(file.filename)[1].lower()

    if file_ext not in [".pdf", ".docx"]:
        logger.warning(
            f"Unsupported file format: {file_ext}",
            extra={"correlation_id": correlation_id},
        )
        raise HTTPException(
            status_code=400,
            detail="Unsupported file format. Please upload a PDF or DOCX file.",
        )

    try:
        # Read the file content
        file_content = await file.read()

        # Extract text based on file type
        if file_ext == ".pdf":
            extracted_text = await extract_text_from_pdf(file_content)
        else:  # .docx
            extracted_text = await extract_text_from_docx(file_content)

        # Check if text was extracted
        if not extracted_text or not extracted_text.strip():
            logger.warning(
                f"No text content extracted from file: {file.filename}",
                extra={"correlation_id": correlation_id},
            )
            return JSONResponse(
                status_code=200,
                content={
                    "status": "warning",
                    "message": "No text content extracted from file",
                    "correlation_id": correlation_id,
                    "filename": file.filename,
                },
            )

        # Create a raw text message
        message = RawTextMessage(
            correlation_id=correlation_id,
            source_system="file_upload",
            source_identifier=file.filename,
            ingestion_timestamp_utc=datetime.utcnow(),
            raw_text=extracted_text,
        )

        # Publish the message to RabbitMQ
        success = rabbitmq_producer.publish_message(message)

        if success:
            logger.info(
                f"File ingestion successful: {file.filename}",
                extra={
                    "correlation_id": correlation_id,
                    "text_length": len(extracted_text),
                },
            )
            return JSONResponse(
                status_code=200,
                content={
                    "status": "success",
                    "message": "File ingested successfully",
                    "correlation_id": correlation_id,
                    "filename": file.filename,
                    "text_length": len(extracted_text),
                },
            )
        else:
            logger.error(
                "Failed to publish message to RabbitMQ",
                extra={"correlation_id": correlation_id},
            )
            raise HTTPException(status_code=500, detail="Failed to process file ingestion")

    except Exception as e:
        logger.error(
            f"Error processing file ingestion: {str(e)}",
            extra={"correlation_id": correlation_id},
        )
        if hasattr(e, "detail"):
            # If the exception has a 'details' attribute, use it for more context
            raise HTTPException(
                status_code=500,
                detail=f"Error processing file ingestion: {str(e.detail)}",
            )
        raise HTTPException(status_code=500, detail=f"Error processing file ingestion: {str(e)}")


@app.get("/health")
async def health_check():
    """Health check endpoint."""
    return {"status": "healthy", "service": "ingestion_file"}


if __name__ == "__main__":
    import uvicorn

    uvicorn.run("main:app", host="0.0.0.0", port=8004, reload=True)
